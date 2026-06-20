"""
Generate the PDF and DOCX files for the knowledge base.
Run this once: python generate_docs.py
"""

from fpdf import FPDF
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

DATA_DIR = os.path.join(os.path.dirname(__file__), "data")


def generate_api_auth_pdf():
    """Generate the API Authentication Guide as a PDF."""
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Page 1 — Title & Overview
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 15, "CloudNova API Authentication Guide", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "This guide covers all authentication methods supported by the CloudNova API, "
        "including API Key authentication, OAuth2 flows, JWT token management, and "
        "troubleshooting common authentication errors."
    )
    pdf.ln(5)

    # Section: API Key Auth
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "1. API Key Authentication", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "API keys are the simplest way to authenticate with the CloudNova API. "
        "Each key is a unique string prefixed with 'cn_api_key_' followed by 40 alphanumeric characters.\n\n"
        "How to generate an API Key:\n"
        "1. Log in to the CloudNova Console\n"
        "2. Navigate to Profile > API Keys > Generate New Key\n"
        "3. Select the scope: read-only, read-write, or admin\n"
        "4. Set an expiration period (90 days recommended)\n"
        "5. Copy the key immediately (it will not be shown again)\n\n"
        "Include the API key in the Authorization header:\n"
        "  Authorization: Bearer cn_api_key_xxxxxxxxxxxxxxxxxxxx\n\n"
        "API Key Best Practices:\n"
        "- Never hardcode API keys in source code\n"
        "- Use environment variables or secret management tools\n"
        "- Rotate keys every 90 days\n"
        "- Use scoped keys with minimum required permissions\n"
        "- Revoke compromised keys immediately at Profile > API Keys > Revoke"
    )
    pdf.ln(5)

    # Section: OAuth2
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "2. OAuth2 Authentication", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "CloudNova supports OAuth2 for applications that need to act on behalf of users.\n\n"
        "Supported Grant Types:\n"
        "- Client Credentials: Server-to-server authentication\n"
        "- Authorization Code: User-delegated access with consent\n"
        "- Authorization Code with PKCE: For mobile and SPA applications\n\n"
        "Client Credentials Flow:\n"
        "1. Register your application at Settings > OAuth Applications\n"
        "2. Obtain client_id and client_secret\n"
        "3. Request a token:\n"
        "   POST https://auth.cloudnova.io/oauth/token\n"
        "   Content-Type: application/x-www-form-urlencoded\n"
        "   grant_type=client_credentials\n"
        "   &client_id=YOUR_CLIENT_ID\n"
        "   &client_secret=YOUR_CLIENT_SECRET\n"
        "   &scope=compute:read compute:write\n\n"
        "4. Use the returned access_token in API requests\n"
        "5. Token expires in 3600 seconds (1 hour)"
    )
    pdf.ln(5)

    # Page 2 — JWT & Errors
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "3. JWT Token Structure", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "CloudNova issues JSON Web Tokens (JWT) for authenticated sessions. "
        "Each token contains three parts: Header, Payload, and Signature.\n\n"
        "Token Payload Fields:\n"
        "- sub: User or service account ID\n"
        "- iss: https://auth.cloudnova.io\n"
        "- aud: https://api.cloudnova.io\n"
        "- exp: Expiration timestamp (Unix epoch)\n"
        "- iat: Issued-at timestamp\n"
        "- scope: Space-separated list of granted scopes\n"
        "- org_id: Organization identifier\n\n"
        "Token Validation:\n"
        "- Verify the signature using CloudNova's public JWKS endpoint\n"
        "- JWKS URL: https://auth.cloudnova.io/.well-known/jwks.json\n"
        "- Check that 'exp' has not passed (tokens are valid for 1 hour)\n"
        "- Verify 'iss' matches https://auth.cloudnova.io\n"
        "- Verify 'aud' matches your API endpoint\n"
        "- System clock skew tolerance: 30 seconds"
    )
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "4. Common Authentication Errors", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Error 401 - INVALID_API_KEY:\n"
        "  Cause: The API key is malformed, expired, or revoked.\n"
        "  Fix: Generate a new API key at Profile > API Keys.\n\n"
        "Error 401 - TOKEN_EXPIRED:\n"
        "  Cause: The JWT access token has expired.\n"
        "  Fix: Request a new token using your refresh token or client credentials.\n"
        "  Note: Tokens expire after 1 hour. Implement automatic token refresh.\n\n"
        "Error 401 - INVALID_SIGNATURE:\n"
        "  Cause: The JWT signature verification failed.\n"
        "  Fix: Ensure you're using the correct JWKS endpoint for verification.\n"
        "  Check that your system clock is synchronized (NTP).\n\n"
        "Error 401 - CLOCK_SKEW:\n"
        "  Cause: Server time differs from your system by more than 30 seconds.\n"
        "  Fix: Synchronize your system clock using NTP.\n"
        "  Linux: sudo ntpdate pool.ntp.org\n"
        "  Windows: w32tm /resync\n\n"
        "Error 403 - INSUFFICIENT_SCOPE:\n"
        "  Cause: The token does not have the required scope for this operation.\n"
        "  Fix: Request a new token with the appropriate scopes.\n"
        "  Available scopes: compute:read, compute:write, storage:read,\n"
        "  storage:write, billing:read, admin:full\n\n"
        "Error 429 - RATE_LIMITED:\n"
        "  Cause: Too many authentication requests.\n"
        "  Fix: Implement exponential backoff. Cache tokens and reuse until expiry.\n"
        "  Auth rate limit: 10 requests per second per client."
    )
    pdf.ln(5)

    # Page 3 — Best Practices
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "5. Security Best Practices for API Authentication", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "1. Use environment variables for credentials:\n"
        "   export CLOUDNOVA_API_KEY=cn_api_key_xxxx\n"
        "   Never commit credentials to version control.\n\n"
        "2. Implement token caching:\n"
        "   Cache access tokens and reuse until close to expiry.\n"
        "   Refresh tokens 5 minutes before expiration.\n\n"
        "3. Use the principle of least privilege:\n"
        "   Request only the scopes your application needs.\n"
        "   Use read-only keys for monitoring and reporting.\n\n"
        "4. Monitor for unauthorized access:\n"
        "   Enable CloudTrail for API authentication events.\n"
        "   Set up alerts for authentication failures.\n"
        "   Review API key usage at Profile > API Keys > Usage.\n\n"
        "5. Implement automatic key rotation:\n"
        "   Use CloudNova's key rotation API to programmatically rotate keys.\n"
        "   POST /api-keys/{key-id}/rotate\n"
        "   The old key remains valid for 24 hours after rotation.\n\n"
        "6. Use IP allowlists for API keys:\n"
        "   Restrict API key usage to specific IP ranges.\n"
        "   Configure at Profile > API Keys > IP Restrictions.\n\n"
        "For additional help, contact support@cloudnova.io or visit\n"
        "https://docs.cloudnova.io/authentication"
    )

    output_path = os.path.join(DATA_DIR, "api_authentication_guide.pdf")
    pdf.output(output_path)
    print(f"Generated: {output_path}")


def generate_refund_policy_docx():
    """Generate the Refund Policy as a DOCX file."""
    doc = Document()

    # Title
    title = doc.add_heading("CloudNova — Refund Policy", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Effective Date: January 1, 2026\n"
        "Last Updated: June 1, 2026"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Section 1
    doc.add_heading("1. Eligibility for Refund", level=1)
    doc.add_paragraph(
        "CloudNova offers refunds under the following circumstances:"
    )
    items = [
        "Service Outage: If CloudNova fails to meet its SLA uptime guarantee (99.99%) "
        "and the customer files a claim within 30 days of the incident.",
        "Billing Error: If a customer is incorrectly charged due to a system error, "
        "a full refund of the overcharged amount will be issued.",
        "Cancellation Within Trial: New customers who cancel within the 14-day free "
        "trial period will not be charged.",
        "Duplicate Charges: If a customer is charged twice for the same service in "
        "a billing cycle, the duplicate charge will be refunded.",
        "Pre-paid Reserved Instances: Customers who cancel reserved instances within "
        "7 days of purchase may receive a prorated refund minus a 10% early termination fee."
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    # Section 2
    doc.add_heading("2. Non-Refundable Items", level=1)
    doc.add_paragraph("The following are NOT eligible for refunds:")
    non_refundable = [
        "On-demand compute usage that has already been consumed",
        "Data transfer charges for outbound data already transferred",
        "Third-party marketplace purchases made through CloudNova",
        "Support tier upgrades (downgrade takes effect next billing cycle)",
        "Domain registration fees",
        "Custom training and consulting services already delivered"
    ]
    for item in non_refundable:
        doc.add_paragraph(item, style="List Bullet")

    # Section 3
    doc.add_heading("3. Refund Process", level=1)
    doc.add_paragraph(
        "To request a refund, follow these steps:"
    )
    steps = [
        "Log in to your CloudNova account and navigate to Support > Billing > Refund Request.",
        "Fill out the refund request form with: your account ID, the billing period in question, "
        "the amount disputed, and the reason for the refund request.",
        "Attach any supporting documentation (invoices, error screenshots, SLA incident reports).",
        "Submit the form. You will receive a confirmation email with a ticket number.",
        "Our billing team will review your request within 5-10 business days.",
        "If approved, the refund will be credited to your original payment method within 15 business days.",
        "For wire transfer refunds, please allow up to 30 business days for processing."
    ]
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"Step {i}: {step}")

    # Section 4
    doc.add_heading("4. SLA Credit vs. Refund", level=1)
    doc.add_paragraph(
        "SLA credits and refunds are different:\n\n"
        "SLA Credits are applied as account credits toward future CloudNova charges. "
        "They are calculated based on the SLA credit schedule (10%, 25%, or 50% of monthly bill) "
        "and do not require full service cancellation.\n\n"
        "Refunds are monetary returns to the original payment method. "
        "Refunds are reserved for billing errors, duplicate charges, and cases where "
        "SLA credits are insufficient to address the customer's concern.\n\n"
        "Customers may request a refund instead of SLA credits for outages exceeding "
        "24 continuous hours, subject to approval by the CloudNova billing team."
    )

    # Section 5
    doc.add_heading("5. Dispute Resolution", level=1)
    doc.add_paragraph(
        "If your refund request is denied and you wish to dispute the decision:\n\n"
        "1. Reply to the refund ticket with additional supporting information.\n"
        "2. Request escalation to a Billing Manager (response within 48 hours).\n"
        "3. If still unresolved, you may request mediation through the CloudNova "
        "Customer Advocacy Program.\n"
        "4. For Enterprise customers, contact your dedicated Technical Account Manager.\n\n"
        "Contact Information:\n"
        "Email: billing@cloudnova.io\n"
        "Phone: +1-800-CLOUDNOVA (option 2 for billing)\n"
        "Hours: Monday-Friday, 9 AM - 6 PM EST"
    )

    # Section 6
    doc.add_heading("6. Policy Changes", level=1)
    doc.add_paragraph(
        "CloudNova reserves the right to modify this refund policy at any time. "
        "Customers will be notified of material changes via email at least 30 days "
        "before the changes take effect. The current version of this policy is always "
        "available at https://www.cloudnova.io/legal/refund-policy."
    )

    output_path = os.path.join(DATA_DIR, "refund_policy.docx")
    doc.save(output_path)
    print(f"Generated: {output_path}")


if __name__ == "__main__":
    os.makedirs(DATA_DIR, exist_ok=True)
    generate_api_auth_pdf()
    generate_refund_policy_docx()
    print("\nAll documents generated successfully!")
